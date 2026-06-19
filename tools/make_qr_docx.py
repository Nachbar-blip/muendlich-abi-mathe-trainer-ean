"""Erzeugt eine .docx mit QR-Code, der auf die Live-App (GitHub Pages) zeigt.
Aufruf: python tools/make_qr_docx.py
Die .docx ist gitignored (lokales Druck-/Teil-Artefakt)."""
import tempfile
from pathlib import Path

import qrcode
from qrcode.constants import ERROR_CORRECT_M
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent
URL = "https://nachbar-blip.github.io/muendlich-abi-mathe-trainer-ean/"
# "hier" = aktueller Arbeitsordner des Nutzers (GAN-Projekt), nicht das EAN-Repo.
OUT = Path("C:/DevProjects/schule/Trainer-Lokal/muendlich-Abi") / "QR Muendlich-Abi Mathe Trainer EAN.docx"

# QR-Bild erzeugen
qr = qrcode.QRCode(version=None, error_correction=ERROR_CORRECT_M, box_size=12, border=2)
qr.add_data(URL)
qr.make(fit=True)
img = qr.make_image(fill_color="black", back_color="white")
tmp_png = Path(tempfile.gettempdir()) / "muendlich_qr.png"
img.save(tmp_png)

# Dokument bauen
doc = Document()

h = doc.add_heading("Mündlich-Abi Mathe — EAN-Trainer", level=0)
h.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Üben für die mündliche Mathe-Prüfung (erhöhtes Niveau): Analysis & Analytische Geometrie")
r.italic = True
r.font.size = Pt(12)

doc.add_paragraph()

# QR zentriert
pq = doc.add_paragraph()
pq.alignment = WD_ALIGN_PARAGRAPH.CENTER
pq.add_run().add_picture(str(tmp_png), width=Cm(8))

cap = doc.add_paragraph()
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
rc = cap.add_run("Handy-Kamera auf den QR-Code halten — oder den Link öffnen:")
rc.font.size = Pt(11)

pl = doc.add_paragraph()
pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
rl = pl.add_run(URL)
rl.bold = True
rl.font.size = Pt(12)
rl.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

doc.add_paragraph()

# Kurzanleitung
b = doc.add_paragraph()
b.alignment = WD_ALIGN_PARAGRAPH.CENTER
rb = b.add_run(
    "Lern-Trainer (Verfahren · Rechnen · Erklären mit Wiederholung) und "
    "Prüfungs-Simulator (20 + 10 + 10 Min). Funktioniert am Handy und am Laptop. "
    "Der Fortschritt wird lokal im Browser gespeichert."
)
rb.font.size = Pt(10)
rb.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.save(str(OUT))
print("QR-Docx erstellt:", OUT)
print("QR-Ziel:", URL)
